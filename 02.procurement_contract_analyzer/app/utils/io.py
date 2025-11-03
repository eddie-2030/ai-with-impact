from typing import Optional
import os
from markdown_it import MarkdownIt
from pdfminer.high_level import extract_text as pdf_extract_text
from docx import Document

def read_text_any(path_or_bytes, filename: Optional[str] = None) -> str:
    # Accept path or bytes-like content; very basic branching
    if isinstance(path_or_bytes, (bytes, bytearray)):
        # try to infer by filename extension
        ext = os.path.splitext(filename or '')[1].lower()
        if ext in ['.md', '.txt']:
            return path_or_bytes.decode('utf-8', errors='ignore')
        if ext == '.pdf':
            # pdfminer expects a file path; we write temp
            tmp = '/tmp/_upload.pdf'
            with open(tmp, 'wb') as f:
                f.write(path_or_bytes)
            return pdf_extract_text(tmp)
        if ext in ['.docx']:
            tmp = '/tmp/_upload.docx'
            with open(tmp, 'wb') as f:
                f.write(path_or_bytes)
            doc = Document(tmp)
            return '\n'.join(p.text for p in doc.paragraphs)
        # fallback
        return path_or_bytes.decode('utf-8', errors='ignore')
    else:
        # treat as file path
        path = str(path_or_bytes)
        ext = os.path.splitext(path)[1].lower()
        if ext == '.pdf':
            return pdf_extract_text(path)
        if ext in ['.docx']:
            doc = Document(path)
            return '\n'.join(p.text for p in doc.paragraphs)
        with open(path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()

def markdown_to_text(md: str) -> str:
    # Cheap MD → text: strip code fences, then rely on parser tokens
    md = md.replace('```', '')
    parser = MarkdownIt()
    tokens = parser.parse(md)
    lines = []
    for t in tokens:
        if t.content:
            lines.append(t.content)
    return '\n'.join(lines)
